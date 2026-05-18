"""
agents/optimizer_agent.py — Optimizer Agent
============================================
CV Tailoring + Cover Letter Generation

FIXED:
  - Groq primary, Gemini fallback (same as analyzer_agent)
  - NO top-level `import streamlit as st` — won't crash agent_runner
  - Streamlit imports only inside show_optimizer_ui() (dashboard only)
  - Added auto-fallback on Groq errors
"""

import os
import io
import re
import logging
from typing import Dict, List
from datetime import datetime

logger = logging.getLogger(__name__)

# ── API Keys ──────────────────────────────────────────────────
GROQ_API_KEY   = os.getenv("GROQ_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
USE_GROQ       = bool(GROQ_API_KEY)


class OptimizerAgent:

    def __init__(self, api_key: str = None):
        self.use_groq = USE_GROQ
        self.gemini_fallback_initialized = False

        if self.use_groq:
            from groq import Groq
            self.groq_client = Groq(api_key=GROQ_API_KEY)
            self.groq_model = "llama-3.3-70b-versatile"
            logger.info("OptimizerAgent using Groq API")

        elif GEMINI_API_KEY or api_key:
            self._init_gemini(api_key)
            self.use_groq = False
            logger.info("OptimizerAgent using Gemini API (fallback)")

        else:
            logger.error("No API key found for OptimizerAgent!")

    def _init_gemini(self, api_key=None):
        """Initialize Gemini client (lazy loading)"""
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key or GEMINI_API_KEY)
            self.gemini_model = genai.GenerativeModel("gemini-2.0-flash")
            self.gemini_fallback_initialized = True
            logger.info("Gemini client initialized for optimizer fallback")
        except Exception as e:
            logger.error(f"Failed to initialize Gemini: {e}")
            self.gemini_fallback_initialized = False

    # ── Public Method (called by agent_runner + dashboard) ────

    def optimize_cv_for_job(
        self,
        cv_text: str,
        jd_text: str,
        job_title: str,
        company_name: str,
        candidate_name: str = "Candidate"
    ) -> Dict:

        keywords = self._extract_keywords(jd_text)
        optimized_cv = self._rewrite_cv(cv_text, jd_text, keywords)
        cover_letter = self._generate_cover_letter(
            candidate_name, job_title, company_name, jd_text, cv_text
        )

        # DOCX only created if python-docx is available
        try:
            docx_bytes = self._create_docx(optimized_cv, candidate_name, job_title, company_name)
        except Exception as e:
            logger.warning(f"DOCX creation skipped: {e}")
            docx_bytes = None

        return {
            "cv_docx": docx_bytes,
            "cv_text": optimized_cv,
            "cover_letter": cover_letter,
            "improvement_score": self._calculate_improvement(cv_text, optimized_cv, keywords),
            "keywords_injected": keywords[:10]
        }

    # ── LLM Calls (WITH AUTO-FALLBACK) ────────────────────────

    def _call_llm(self, prompt: str, max_tokens: int = 1000) -> str:
        """Call LLM with auto-fallback from Groq to Gemini"""
        if self.use_groq:
            try:
                return self._call_groq(prompt, max_tokens)
            except Exception as groq_error:
                logger.warning(f"Groq failed, switching to Gemini: {groq_error}")
                if GEMINI_API_KEY:
                    if not self.gemini_fallback_initialized:
                        self._init_gemini()
                    return self._call_gemini(prompt, max_tokens)
                raise
        else:
            return self._call_gemini(prompt, max_tokens)

    def _call_groq(self, prompt: str, max_tokens: int = 1000) -> str:
        try:
            response = self.groq_client.chat.completions.create(
                model=self.groq_model,
                messages=[
                    {"role": "system", "content": "You are an expert CV writer and career coach."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,
                temperature=0.4
            )
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Groq optimizer error: {e}")
            raise

    def _call_gemini(self, prompt: str, max_tokens: int = 1000) -> str:
        response = self.gemini_model.generate_content(
            prompt,
            generation_config={"temperature": 0.4, "max_output_tokens": max_tokens}
        )
        return response.text

    # ── CV Rewriting ──────────────────────────────────────────

    def _rewrite_cv(self, cv_text: str, jd_text: str, keywords: List[str]) -> str:
        prompt = f"""
Rewrite this CV to be ATS-friendly for the following job.

=== ORIGINAL CV ===
{cv_text[:2000]}

=== TARGET KEYWORDS ===
{', '.join(keywords[:15])}

=== RULES ===
- Keep the same structure and sections
- Inject keywords naturally (don't force them)
- Use strong action verbs (Built, Developed, Optimized, Led)
- Add quantified results where possible (%, numbers)
- Do NOT invent experience the candidate doesn't have

Return ONLY the rewritten CV text, no explanation.
"""
        try:
            return self._call_llm(prompt, max_tokens=1500)
        except Exception as e:
            logger.warning(f"CV rewrite failed, returning original: {e}")
            return cv_text

    def _extract_keywords(self, jd_text: str) -> List[str]:
        prompt = f"""
Extract the top 15 most important technical skills and requirements from this job description.

{jd_text[:1500]}

Return ONLY a comma-separated list. Example:
Python, TensorFlow, REST APIs, Docker, SQL
"""
        try:
            result = self._call_llm(prompt, max_tokens=200)
            return [k.strip() for k in result.split(',')][:15]
        except Exception as e:
            logger.warning(f"Keyword extraction failed: {e}")
            return ["Python", "Machine Learning", "Communication"]

    # ── Cover Letter ──────────────────────────────────────────

    def _generate_cover_letter(
        self,
        name: str,
        job_title: str,
        company: str,
        jd_text: str,
        cv_text: str
    ) -> str:
        today = datetime.now().strftime('%B %d, %Y')

        prompt = f"""
Write a professional cover letter (150-200 words) for {name} applying for {job_title} at {company}.

=== JOB DESCRIPTION ===
{jd_text[:800]}

=== CANDIDATE BACKGROUND ===
{cv_text[:600]}

=== STRUCTURE ===
Paragraph 1: Enthusiasm for {company} and this role (3-4 sentences)
Paragraph 2: 2-3 specific matching skills with brief examples (4-5 sentences)  
Paragraph 3: Why great fit + request for interview (2-3 sentences)

=== REQUIREMENTS ===
- 150-200 words total
- Professional but warm tone
- No bullet points, no markdown
- Use specific details from the candidate's background

Return ONLY the letter body (no subject line, no "Dear" header — just paragraphs):
"""
        try:
            body = self._call_llm(prompt, max_tokens=500)

            # Check minimum length
            if len(body.split()) < 80:
                return self._fallback_cover_letter(name, job_title, company, today)

            # Wrap with proper header/footer
            return f"""{today}

Hiring Manager
{company}

Dear Hiring Manager,

{body.strip()}

Sincerely,
{name}"""

        except Exception as e:
            logger.error(f"Cover letter generation error: {e}")
            return self._fallback_cover_letter(name, job_title, company, today)

    def _fallback_cover_letter(self, name: str, job_title: str, company: str, today: str) -> str:
        return f"""{today}

Hiring Manager
{company}

Dear Hiring Manager,

I am writing to apply for the {job_title} position at {company}. With my background in Python development, machine learning, and AI systems, I am excited about the opportunity to contribute to your team.

My experience includes building and deploying end-to-end ML pipelines, developing REST APIs, and working with cloud platforms. I have delivered multiple AI projects from concept to production, and I am confident I can bring similar results to {company}.

I would welcome the opportunity to discuss how my skills align with your needs. Thank you for your time and consideration.

Sincerely,
{name}"""

    # ── DOCX Creation ─────────────────────────────────────────

    def _create_docx(self, cv_text: str, name: str, job_title: str, company: str) -> bytes:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        display_name = name if (name and name != "Candidate") else "Candidate"

        title = doc.add_heading(display_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run("📍 Location | 📧 email@example.com | 🔗 linkedin.com/in/profile")

        doc.add_paragraph()

        for line in cv_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('##') or (line.isupper() and len(line) < 30):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith(('-', '•', '*')):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def _calculate_improvement(self, original: str, optimized: str, keywords: List[str]) -> int:
        orig_lower = original.lower()
        opt_lower = optimized.lower()
        new_kw = [k for k in keywords if k.lower() in opt_lower and k.lower() not in orig_lower]
        return min(100, int((len(new_kw) / max(len(keywords), 1)) * 100))


# ═══════════════════════════════════════════════════════════════
# STREAMLIT UI — Only called from dashboard pages, NOT agent_runner
# ═══════════════════════════════════════════════════════════════

def show_optimizer_ui(job: Dict, db):
    import streamlit as st  # ✅ Local import — safe, won't crash agent_runner

    st.markdown("---")
    st.markdown("### ✏️ Tailor CV for this Job")

    col1, col2 = st.columns(2)

    with col1:
        if st.button("📄 Generate Tailored CV", key=f"optimize_{job['id']}"):
            cv_text = st.session_state.get("cv_text", "")
            if not cv_text:
                st.warning("⚠️ Pehle Settings mein CV upload karo!")
                return

            with st.spinner("🤖 Generating tailored CV and cover letter... (10-15 seconds)"):
                optimizer = OptimizerAgent()
                result = optimizer.optimize_cv_for_job(
                    cv_text=cv_text,
                    jd_text=job.get("description", ""),
                    job_title=job.get("title", ""),
                    company_name=job.get("company", ""),
                    candidate_name=st.session_state.get("user_name", "Candidate")
                )
                st.session_state[f"optimized_cv_{job['id']}"] = result
                st.success(f"✅ CV optimized! {result['improvement_score']}% improvement")
                st.rerun()

    with col2:
        result = st.session_state.get(f"optimized_cv_{job['id']}")
        if result and result.get("cv_docx"):
            filename = re.sub(
                r'[<>:"/\\|?*]', '_',
                f"{job.get('company', 'Job')}_{job.get('title', 'Role')}_CV.docx"
            )
            st.download_button(
                label="⬇️ Download Tailored CV (DOCX)",
                data=result["cv_docx"],
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_{job['id']}"
            )

    result = st.session_state.get(f"optimized_cv_{job['id']}")
    if result:
        st.markdown("---")
        st.markdown("### 📝 Cover Letter")

        cover_text = result.get("cover_letter", "No cover letter generated")
        st.text_area(
            label="Cover Letter",
            value=cover_text,
            height=400,
            key=f"cover_{job['id']}",
            label_visibility="collapsed"
        )

        word_count = len(cover_text.split())
        if word_count < 100:
            st.warning(f"⚠️ Cover letter is {word_count} words (recommended: 150-200).")
        else:
            st.caption(f"📝 {word_count} words")

        keywords = result.get("keywords_injected", [])
        if keywords:
            st.caption(f"✨ Keywords added: {', '.join(keywords[:8])}")
